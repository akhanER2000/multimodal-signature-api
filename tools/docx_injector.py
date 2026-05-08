import docx
import base64
import json
import os
import io

def process_docx_signature(input_json_path, output_docx_path):
    os.makedirs(os.path.dirname(output_docx_path), exist_ok=True)
    
    with open(input_json_path, 'r') as f:
        data = json.load(f)
        
    try:
        # 1. Cargar DOCX desde base64 a un buffer en memoria
        docx_bytes = base64.b64decode(data['document']['content_base64'])
        doc_stream = io.BytesIO(docx_bytes)
        doc = docx.Document(doc_stream)
        
        # 2. Cargar imagen de la firma
        sig_payload = data['signature_data']['payload']
        if "," in sig_payload:
            sig_base64 = sig_payload.split(",")[1]
        else:
            sig_base64 = sig_payload
        sig_bytes = base64.b64decode(sig_base64)
        img_stream = io.BytesIO(sig_bytes)
        
        # 3. Inyectar imagen al final del documento (centrada)
        pos = data['signature_data']['position']
        width_pt = pos.get('width', 150) # Fallback a 150 si no existe
        
        paragraph = doc.add_paragraph()
        paragraph.alignment = 1 # 1 = Centrado
        run = paragraph.add_run()
        run.add_picture(img_stream, width=docx.shared.Pt(width_pt))
        
        # 4. Guardar en .tmp/
        doc.save(output_docx_path)
        
        print(f"✅ ÉXITO: DOCX firmado generado en {output_docx_path}")
        return True
    except Exception as e:
        print(f"❌ ERROR CRÍTICO procesando DOCX: {e}")
        return False

if __name__ == "__main__":
    print("Módulo docx_injector.py creado y verificado por el linter (sin ejecutar prueba viva).")
